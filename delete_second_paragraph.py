import docx

doc = docx.Document('/root/.openclaw/qqbot/downloads/1_1776043418375.docx')

if len(doc.paragraphs) > 1:
    doc.paragraphs[1].text = ''

doc.save('/root/.openclaw/qqbot/downloads/1_1776043418375_修改后.docx')