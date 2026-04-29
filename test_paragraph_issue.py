#!/usr/bin/env python3
"""
演示 word-docx 技能可能存在的段落处理问题
"""

from docx import Document

def demonstrate_paragraph_issue():
    """演示段落处理问题"""
    
    print("=" * 60)
    print("演示 word-docx 段落处理问题")
    print("=" * 60)
    
    # 创建一个测试文档
    doc = Document()
    
    # 添加一些文本（模拟用户输入）
    test_text = """对memory-pkm的改进

对重要对话（对话≥5 轮或含关键信息）结束后马上归档，对conversations设置提炼到knowledge后即可清理。

对word-docx的改进。

按换行符 \\n来区分段落。

对ocr-python的改进。

提取没有文字的图片时返回未识别到文字。"""
    
    print("原始文本:")
    print(test_text)
    print("\n" + "=" * 60)
    
    # 问题演示1: 不同的添加文本方式
    print("\n问题1: 不同的文本添加方式产生不同的段落结构")
    print("-" * 40)
    
    # 方式A: 直接添加整个文本（可能有问题）
    doc_a = Document()
    doc_a.add_paragraph(test_text)
    print("方式A (直接添加整个文本):")
    print(f"  段落数量: {len(doc_a.paragraphs)}")
    print(f"  段落1内容: '{doc_a.paragraphs[0].text[:50]}...'")
    
    # 方式B: 按行分割添加（正确方式）
    doc_b = Document()
    lines = test_text.split('\n')
    for line in lines:
        if line.strip():  # 跳过空行
            doc_b.add_paragraph(line.strip())
    print("\n方式B (按行分割添加):")
    print(f"  段落数量: {len(doc_b.paragraphs)}")
    for i, para in enumerate(doc_b.paragraphs[:3], 1):
        print(f"  段落{i}内容: '{para.text}'")
    
    print("\n" + "=" * 60)
    
    # 问题演示2: 文本提取时的段落处理
    print("\n问题2: 文本提取时的段落处理不一致")
    print("-" * 40)
    
    # 创建一个有多个段落的文档
    doc_c = Document()
    doc_c.add_paragraph("第一段内容")
    doc_c.add_paragraph("")  # 空段落
    doc_c.add_paragraph("第二段内容")
    doc_c.add_paragraph("第三段内容，\n包含换行符")
    
    print("文档C结构:")
    for i, para in enumerate(doc_c.paragraphs, 1):
        print(f"  段落{i}: '{para.text}' (长度: {len(para.text)})")
    
    # 不同的提取方式
    print("\n不同的文本提取方式:")
    
    # 方式1: 简单连接（可能有问题）
    text1 = ""
    for para in doc_c.paragraphs:
        text1 += para.text + "\n"
    print("方式1 (简单连接):")
    print(f"  提取的文本: {repr(text1)}")
    
    # 方式2: 过滤空段落
    text2 = ""
    for para in doc_c.paragraphs:
        if para.text.strip():
            text2 += para.text + "\n"
    print("\n方式2 (过滤空段落):")
    print(f"  提取的文本: {repr(text2)}")
    
    # 方式3: 使用双换行符分隔
    paragraphs = [para.text.strip() for para in doc_c.paragraphs if para.text.strip()]
    text3 = "\n\n".join(paragraphs)
    print("\n方式3 (双换行符分隔):")
    print(f"  提取的文本: {repr(text3)}")
    
    print("\n" + "=" * 60)
    
    # 建议的解决方案
    print("\n建议的解决方案:")
    print("-" * 40)
    
    class DocxParagraphProcessor:
        """段落处理器"""
        
        def __init__(self, paragraph_delimiter='\n\n'):
            self.paragraph_delimiter = paragraph_delimiter
        
        def text_to_document(self, text, doc=None):
            """将文本转换为 Word 文档"""
            if doc is None:
                doc = Document()
            
            # 按照双换行符分割段落
            paragraphs = self._split_paragraphs(text)
            
            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())
            
            return doc
        
        def document_to_text(self, doc):
            """将 Word 文档转换为文本"""
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            return self.paragraph_delimiter.join(paragraphs)
        
        def _split_paragraphs(self, text):
            """分割文本为段落"""
            # 规范化换行符
            import re
            normalized = re.sub(r'\n\s*\n+', self.paragraph_delimiter, text.strip())
            paragraphs = [p.strip() for p in normalized.split(self.paragraph_delimiter) if p.strip()]
            return paragraphs
    
    # 使用建议的解决方案
    processor = DocxParagraphProcessor()
    
    print("使用建议的段落处理器:")
    print(f"  原始文本段落数: {len(processor._split_paragraphs(test_text))}")
    
    # 转换为文档
    doc_processed = processor.text_to_document(test_text)
    print(f"  生成的文档段落数: {len(doc_processed.paragraphs)}")
    
    # 转换回文本
    text_processed = processor.document_to_text(doc_processed)
    print(f"  转换回的文本: {repr(text_processed[:100])}...")

if __name__ == "__main__":
    demonstrate_paragraph_issue()